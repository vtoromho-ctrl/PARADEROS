# report_generator.py

# --- IMPORTACIONES ---
import io
from datetime import datetime #
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from googleapiclient.http import MediaIoBaseDownload
import pandas as pd
from typing import List, Any, Dict
import numpy as np

# ===================================================================
# FUNCIONES DE ESTILO Y FORMATO
# ===================================================================
document = Document()
LOGO_DRIVE_FILE_ID = None
IMG_UBICACION_PROYECTO_ID = None
IMG_UBICACION_PARADAS_ID = None


def definir_estilos_base(document):

    # Define una función interna para no repetir código
    def aplicar_estilo(nombre_estilo, fuente, tamano, negrita, alineacion=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
        style = document.styles[nombre_estilo]
        font = style.font
        font.name = fuente
        font.size = Pt(tamano)
        font.bold = negrita
        font.color.rgb = RGBColor(0, 0, 0)
        p_fmt = style.paragraph_format
        p_fmt.alignment = alineacion
        # Asegura compatibilidad con fuentes de Asia Oriental
        rpr = style.element.rPr
        rpr.rFonts.set(qn('w:eastAsia'), fuente)

    # Aplicamos los estilos a los títulos que usaremos
    aplicar_estilo("Heading 1", "Arial Narrow", 11, True)
    aplicar_estilo("Heading 2", "Arial Narrow", 11, True)
    aplicar_estilo("Heading 3", "Arial Narrow", 11, True,alineacion=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    aplicar_estilo("Heading 4", "Arial Narrow", 9, True,alineacion=WD_PARAGRAPH_ALIGNMENT.CENTER)

def agregar_imagen_con_formato_drive(document, service_drive, file_id, descripcion, estado, fuente="Fuente: Elaboración propia."):
    print(f"   - Agregando imagen con formato: {descripcion}")
    capitulo = estado["capitulo"]
    num_figura = estado["figura"]
    etiqueta = f"Figura {capitulo}.{num_figura}. {descripcion}"

    # --- Título de la figura ---
    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_titulo.paragraph_format.space_before = Pt(0) 
    p_titulo.paragraph_format.space_after  = Pt(0)  
    run_titulo = p_titulo.add_run(etiqueta)
    run_titulo.font.name = "Arial Narrow"
    run_titulo.font.size = Pt(9)
    run_titulo.bold = True

    # --- Imagen ---
    try:
        request = service_drive.files().get_media(fileId=file_id)
        file_bytes = io.BytesIO()
        downloader = MediaIoBaseDownload(file_bytes, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        file_bytes.seek(0)

        p_img = document.add_paragraph()
        p_img.paragraph_format.space_before = Pt(0) 
        p_img.paragraph_format.space_after  = Pt(0)
        p_img.add_run().add_picture(file_bytes, width=Inches(5.3))
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        document.add_paragraph(f"[Error al cargar imagen ID: {file_id}]")
        print(f"   ✗ Error: No se pudo agregar la imagen {file_id}. Error: {e}")

    # --- Fuente de la figura ---
    p_fuente = document.add_paragraph()
    p_fuente.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_fuente = p_fuente.add_run(fuente)
    run_fuente.font.name = "Arial Narrow"
    run_fuente.font.size = Pt(9)
    run_fuente.italic = False
    run_fuente.bold = True

    estado["figura"] += 1 # Incrementar contador para la siguiente figura

def agregar_imagen_simple_drive(document, service_drive, file_id, width_inch=6.0, paragraph=None):
    try:
        request = service_drive.files().get_media(fileId=file_id)
        file_bytes = io.BytesIO()
        downloader = MediaIoBaseDownload(file_bytes, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        file_bytes.seek(0)

        # Si no se nos da un párrafo, creamos uno nuevo en el documento.
        if paragraph is None:
            paragraph = document.add_paragraph()

        # Añadimos la imagen al párrafo correspondiente.
        paragraph.add_run().add_picture(file_bytes, width=Inches(width_inch))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        print(f"   ✓ Imagen simple {file_id} agregada.")
    except Exception as e:
        print(f"   ✗ Advertencia: No se pudo agregar la imagen simple {file_id}. Error: {e}")
def aplicar_color_celda(celda, color_hex="D9D9D9"):
    """Pone un color de fondo a una celda de la tabla."""
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def agregar_tabla_formateada(
    document,
    descripcion: str,
    estado: Dict[str, Any],
    fuente: str = "Fuente: MHO Consultores",
    tabla_data: List[Dict[str, Any]] = None,
    headers: List[str] = None,
    rows: List[Any] = None,
):
    """
    Crea una tabla con formato avanzado en el documento.
    Acepta datos en dos formas:
      - tabla_data: lista de dicts con claves (caracteristica/cumplimiento/observacion)
      - headers/rows: encabezados + filas (listas o dicts)

    Requisitos:
      - estado: dict con 'capitulo' y 'cuadro' (se incrementa 'cuadro' al final)
    """

    # -------------------- Etiqueta y título --------------------
    capitulo = estado.get("capitulo", 0)
    num_cuadro = estado.get("cuadro", 1)
    etiqueta = f"Cuadro {capitulo}.{num_cuadro}. {descripcion.strip()}"

    p_titulo = document.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(0)  # sin espacio
    p_titulo.paragraph_format.space_after  = Pt(0)  # sin espacio
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_titulo = p_titulo.add_run(etiqueta)
    run_titulo.font.name = "Arial Narrow"
    run_titulo.font.size = Pt(9)
    run_titulo.bold = True

    # -------------------- Construcción del DataFrame --------------------
    # Caso A: tabla_data (lista de dicts)
    if tabla_data is not None:
        if isinstance(tabla_data, dict):
            tabla_data = [tabla_data]
        df = pd.DataFrame(tabla_data)
        if not df.empty:
            # normaliza y renombra columnas
            df.columns = [str(c).strip().lower() for c in df.columns]
            rename_map = {
                'caracteristica': 'Característica',
                'cumplimiento':  'Cumplimiento',
                'observacion':   'Observación',
            }
            df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
            order = [c for c in ['Característica', 'Cumplimiento', 'Observación'] if c in df.columns]
            if order:
                df = df[order]
        else:
            df = pd.DataFrame(columns=['Característica', 'Cumplimiento', 'Observación'])

    # Caso B: headers/rows
    else:
        cols: List[str] = [str(c) for c in (headers or [])]
        rows_list: List[List[Any]] = []
        for r in (rows or []):
            if isinstance(r, dict):
                row_vals = [r.get(h, "") for h in cols]
            else:
                row_vals = list(r)
                # Alinea longitud de fila a columnas
                if len(row_vals) < len(cols):
                    row_vals += [""] * (len(cols) - len(row_vals))
                elif len(row_vals) > len(cols):
                    row_vals = row_vals[:len(cols)]
            # Limpia None
            row_vals = ["" if v is None else v for v in row_vals]
            rows_list.append(row_vals)

        df = pd.DataFrame.from_records(rows_list, columns=cols)

        # Si no vino nada, crea DF vacío con columnas estándar
        if df.empty:
            df = pd.DataFrame(columns=['Característica', 'Cumplimiento', 'Observación'])

    # Si sigue vacío y sin columnas, crea estructura mínima
    if df.empty and df.columns.size == 0:
        df = pd.DataFrame(columns=['Característica', 'Cumplimiento', 'Observación'])

    # -------------------- Crear tabla Word --------------------
    ncols = len(df.columns)
    if ncols == 0:
        # Evita tabla sin columnas
        df = pd.DataFrame(columns=['Característica', 'Cumplimiento', 'Observación'])
        ncols = 3

    table = document.add_table(rows=1, cols=ncols)
    table.style = 'Table Grid'  # usa tu estilo si tienes uno definido

    # --- Encabezados ---
    hdr_cells = table.rows[0].cells
    for j, col_name in enumerate(df.columns):
        p = hdr_cells[j].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)   # sin espacios extra
        p.paragraph_format.space_after  = Pt(0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(str(col_name))
        run.font.name = "Arial Narrow"
        run.font.size = Pt(9)
        run.bold = True
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        try:
            aplicar_color_celda(hdr_cells[j])  # si tienes la función utilitaria
        except NameError:
            pass

    # --- Filas ---
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row_data):
            p = row_cells[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)   # sin espacios extra
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("" if (pd.isna(val)) else str(val))
            run.font.name = "Arial Narrow"
            run.font.size = Pt(9)
            run.bold = False
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # -------------------- Fuente --------------------
    p_fuente = document.add_paragraph()
    p_fuente.paragraph_format.space_before = Pt(0)  # sin espacio
    p_fuente.paragraph_format.space_after  = Pt(0)  # sin espacio
    p_fuente.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_fuente = p_fuente.add_run(fuente)
    run_fuente.font.name = "Arial Narrow"
    run_fuente.font.size = Pt(8)
    run_fuente.bold = True

    # -------------------- Actualiza estado --------------------
    estado["cuadro"] = int(estado.get("cuadro", 1)) + 1

    # Log útil para depurar
    print("   ✓ Tabla formateada creada. Filas:", len(df.index), "| Cols:", list(df.columns))
                                                                                 
def agregar_tabla_desde_df(document, df, descripcion, estado, fuente, **kwargs):
    """
    Función adaptada de tu ejemplo para crear una tabla directamente desde un DataFrame.
    """
    # Esta función sería una adaptación de tu función 'agregar_tabla' original
    # para asegurar que maneja el formato, colores, y combinaciones que necesitas.
    # Por simplicidad aquí usamos una versión más directa.

    print(f"   - Creando tabla desde DataFrame: {descripcion}")
    capitulo = estado["capitulo"]
    num_cuadro = estado["cuadro"]
    etiqueta = f"Cuadro {capitulo}.{num_cuadro}. {descripcion.strip()}"

    # Título de la tabla
    agregar_texto(document, etiqueta).runs[0].bold = True

    # Crear tabla
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Contenido
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val if not pd.isna(val) else "")

    # Fuente
    document.add_paragraph(fuente).runs[0].font.size = Pt(8)

    estado["cuadro"] += 1
def read_excel_from_drive(service, file_id):
    """Descarga un archivo Excel de Drive y lo carga en un DataFrame de Pandas."""
    if not file_id: return None
    try:
        request = service.files().get_media(fileId=file_id)
        file_bytes = io.BytesIO()
        downloader = MediaIoBaseDownload(file_bytes, request)
        # ... (lógica del downloader)
        file_bytes.seek(0)
        return pd.read_excel(file_bytes)
    except Exception as e:
        print(f"Error al leer Excel desde Drive (ID: {file_id}): {e}")
        return None

def crear_tabla_evidencia(document, service_drive, titulo_tabla, seccion_analisis):
    """
    Crea una tabla de 1 columna para mostrar la evidencia fotográfica (título, imagen, descripción).
    """
    print(f"   - Creando tabla de evidencia: {titulo_tabla}")

    # Extraer datos de la sección
    image_ids = seccion_analisis.get("image_ids", [])
    description = seccion_analisis.get("description", "No hay descripción disponible.")

    if not image_ids:
        return # No crear la tabla si no hay imágenes

    # Crear tabla de 1 columna
    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(6.5)

    # Fila 1: Título con color
    cell_titulo = table.rows[0].cells[0]
    cell_titulo.text = ''
    cell_titulo.paragraphs[0].runs[0].bold = True
    p_titulo = cell_titulo.add_paragraph(titulo_tabla, style="Heading 3")
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    aplicar_color_celda(cell_titulo) # Asumiendo que tienes la función aplicar_color_celda

    # Filas siguientes: Una para cada imagen
    for img_id in image_ids:
        cell_img = table.add_row().cells[0]
        p_img = cell_img.paragraphs[0]
        # Usamos una función simple para agregar la imagen sin texto adicional
        agregar_imagen_simple_drive(document=None, paragraph=p_img, service_drive=service_drive, file_id=img_id, width_inch=6.0)
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Fila final: Descripción
    cell_desc = table.add_row().cells[0]
    cell_desc.text = '' # 1. Limpia el párrafo por defecto
    p_desc = cell_desc.add_paragraph(description, style="Normal")
    p_desc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def cambiar_capitulo(estado, nuevo_capitulo):
    estado["capitulo"] = nuevo_capitulo
    estado["figura"] = 1
    estado["cuadro"] = 1

def agregar_titulo(document, texto, estilo="Heading 1"):
    p = document.add_paragraph(style=estilo)
    run = p.add_run(texto)
    run.font.name = "Arial Narrow"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial Narrow")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p

def agregar_subtitulo(document, texto, estilo="Heading 2"):
    p = document.add_paragraph(style=estilo)
    run = p.add_run(texto)
    run.font.name = "Arial Narrow"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial Narrow")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p

def agregar_subsub(document, texto, estilo="Heading 3"):
    p = document.add_paragraph(style=estilo)
    run = p.add_run(texto)
    run.font.name = "Arial Narrow"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial Narrow")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p

def agregar_texto(document, texto, estilo="Normal"):
    p = document.add_paragraph(style=estilo)
    run = p.add_run(texto)
    run.font.name = "Arial Narrow"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial Narrow")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run.italic = False
    return p

def agregar_lista(document, items, estilo='bullet', fuente="Arial Narrow", tamano=11, interlineado=Pt(12)):

    estilo_word = 'List Bullet' if estilo == 'bullet' else 'List Number'

    for item in items:
        p = document.add_paragraph(style=estilo_word)
        run = p.add_run(item)
        run.font.name = fuente
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fuente)
        run.font.size = Pt(tamano)
        p.paragraph_format.line_spacing = interlineado
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def agregar_espacio(document, cantidad=1):
    for _ in range(cantidad):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
# ===================================================================
# FUNCIÓN PRINCIPAL PARA CREAR EL INFORME
# ===================================================================

def crear_informe_paraderos(datos_informe,
        service_drive,
        drive_file_ids=None,
        logo_id=None,
        tablas_id=None,
        img_ubicacion_proyecto_id=None,
        img_ubicacion_paradas_id=None,
    ):
    try:
        print("🚀 Iniciando la generación del informe...")
        document = Document()
        estado_informe = {"capitulo": 1, "figura": 1, "cuadro": 1}
        definir_estilos_base(document)

        # --- DATOS DEL PROYECTO ---
        
        # Normaliza: prioriza el dict si viene completo; si no, usa los args sueltos
        drive_file_ids = drive_file_ids or {}
        logo_id = drive_file_ids.get("logo_id") or logo_id
        tablas_id = drive_file_ids.get("tablas_id") or tablas_id
        img_ubicacion_proyecto_id = drive_file_ids.get("img_ubicacion_proyecto_id") or img_ubicacion_proyecto_id
        img_ubicacion_paradas_id = drive_file_ids.get("img_ubicacion_paradas_id") or img_ubicacion_paradas_id

        drive_ids = datos_informe.get("drive_file_ids", {})
        logo_id = drive_ids.get("logo")
        tablas_id = drive_ids.get("tablas")
        img_ubicacion_proyecto_id = drive_ids.get("ubicacion_proyecto")
        img_ubicacion_paradas_id = drive_ids.get("ubicacion_paradas")
        info_proyecto = datos_informe.get("info_proyecto", {})
        nombre_proyecto = info_proyecto.get("proyecto", "[Nombre del Proyecto]")
        comuna = info_proyecto.get("comuna", "[Comuna]")
        estudio = info_proyecto.get("estudio", "[Tipo de Estudio]")
        mitigacion = info_proyecto.get("mitigacion", "[N° Mitigación]")
        resolucion = info_proyecto.get("resolucion", "[N° Resolución]")
        fecha = info_proyecto.get("fecha", "[Fecha Resolución]")
        medida_mitigacion = info_proyecto.get("medida_mitigacion", "[Descripción de la Medida]")
        ubi_proyecto = info_proyecto.get("ubi_proyecto", "[Ubicación del Proyecto]")
        region = info_proyecto.get("region", "[Región]")

        # ==========================================================
        # PORTADA (Lógica integrada de tu ejemplo)
        # ==========================================================
        print("   - Creando portada...")

        # Espaciado vertical inicial
        for _ in range(6): document.add_paragraph()

        # Título principal
        titulo = document.add_paragraph()
        titulo_run = titulo.add_run(
            f"MEJORAMIENTO DE PARADAS DE TRANSPORTE PÚBLICO\n"
            f"MEDIDA DE MITIGACIÓN {(info_proyecto.get('mitigacion') or '').upper()} {(info_proyecto.get('estudio') or '').upper()}\n"
            f"{(info_proyecto.get('proyecto') or '').upper()}\n"
            f"{(info_proyecto.get('comuna') or '').upper()}\n"
        )
        titulo_run.bold = True
        titulo_run.font.size = Pt(11)
        titulo_run.font.name = "Arial Narrow"
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titulo.paragraph_format.line_spacing = Pt(22)

        # Espaciado y fecha automática
        for _ in range(8): document.add_paragraph()
        fecha_actual = datetime.now().strftime("%B %Y").upper() # Genera "AGOSTO 2025"
        p_fecha = document.add_paragraph()
        p_fecha.add_run(fecha_actual).bold = True
        p_fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Información de la empresa
        for _ in range(2): document.add_paragraph()
        p_empresa = document.add_paragraph()
        p_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_empresa = p_empresa.add_run("Avenida Presidente Riesco 5335 Oficina 606, Las Condes \nTeléfono: (56 2) 2 657 16 25\ncontacto@mho.cl - www.mho.cl")
        run_empresa.font.name = "Arial Narrow"
        run_empresa.font.size = Pt(8)

        document.add_page_break()

        # --- ÍNDICE (Placeholder) ---
        agregar_titulo(document, "ÍNDICE GENERAL")
        print("   ✓ Portada e Índice creados.")
        section_main = document.add_section(WD_SECTION.NEW_PAGE)

        # ==========================================================
            # ENCABEZADO PERSONALIZADO (Tu código integrado)
            # ==========================================================
        print("   - Creando encabezado...")
        header = section_main.header
        header.is_linked_to_previous = False

        table_h = header.add_table(rows=1, cols=2, width=Inches(6.5))

        # Borde inferior para la tabla del encabezado
        tbl_h = table_h._tbl
        tblPr_h = tbl_h.tblPr
        if tblPr_h is None:
            tblPr_h = OxmlElement('w:tblPr')
            tbl_h.insert(0, tblPr_h)
        tblBorders_h = OxmlElement('w:tblBorders')
        bottom_h = OxmlElement('w:bottom')
        bottom_h.set(qn('w:val'), 'single')
        bottom_h.set(qn('w:sz'), '6')
        bottom_h.set(qn('w:color'), '000000')
        tblBorders_h.append(bottom_h)
        tblPr_h.append(tblBorders_h)

        # Celda izquierda: nombre del proyecto
        cell1_h = table_h.cell(0, 0)
        p1_h = cell1_h.paragraphs[0]
        p1_h.add_run(f"Informe de Paradero - {nombre_proyecto}").font.size = Pt(8)

        # Celda derecha: número de página
        cell2_h = table_h.cell(0, 1)
        p2_h = cell2_h.paragraphs[0]
        p2_h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run2_h = p2_h.add_run()
        run2_h.font.size = Pt(8)

        # Añadir campo de NÚMERO DE PÁGINA
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run2_h._r.append(fldChar1); run2_h._r.append(instrText); run2_h._r.append(fldChar2)

        # ==========================================================
        # PIE DE PÁGINA PERSONALIZADO
        # ==========================================================
        print("   - Creando pie de página...")
        footer = section_main.footer
        footer.is_linked_to_previous = False

        table_f = footer.add_table(rows=1, cols=2, width=Inches(6.5))

        # Borde superior para la tabla del pie de página
        tbl_f = table_f._tbl
        tblPr_f = tbl_f.tblPr
        if tblPr_f is None:
            tblPr_f = OxmlElement('w:tblPr')
            tbl_f.insert(0, tblPr_f)
        tblBorders_f = OxmlElement('w:tblBorders')
        top_f = OxmlElement('w:top')
        top_f.set(qn('w:val'), 'single')
        top_f.set(qn('w:sz'), '6')
        top_f.set(qn('w:color'), '000000')
        tblBorders_f.append(top_f)
        tblPr_f.append(tblBorders_f)

        # Celda izquierda: Logo desde Google Drive
        cell_img = table_f.cell(0, 0)
        p_img = cell_img.paragraphs[0]
        if logo_id: agregar_imagen_simple_drive(document=None, paragraph=p_img, service_drive=service_drive, file_id=logo_id, width_inch=0.65)
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Celda derecha: texto de contacto
        cell_txt = table_f.cell(0, 1)
        texto_footer = "Presidente Riesco 5335 Of. 606, Las Condes \nTeléfono: (56 2) 2 657 1625 \ncontacto@mho.cl \nwww.mho.cl"
        p_txt = cell_txt.paragraphs[0]
        p_txt.add_run(texto_footer).font.size = Pt(7)
        p_txt.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p_txt.paragraph_format.line_spacing = Pt(10)

        # ==========================================================
        # CAPÍTULO 1: ANTECEDENTES
        # ==========================================================
        print("   - Creando Capítulo 1: Antecedentes...")
        cambiar_capitulo(estado_informe, 1)
        agregar_titulo(document, "1. ANTECEDENTES")
        agregar_espacio(document)
        agregar_texto(document, f"El presente estudio, tiene por objetivo dar cumplimiento a la medida de mitigación {info_proyecto.get('mitigacion', '[mitigacion]')} del {info_proyecto.get('estudio', '[estudio]')} aprobado para {info_proyecto.get('proyecto', '[proyecto]')}. Las mitigaciones que se abordan a continuación tienen relación con el mantenimiento y reparación de la infraestructura y elementos de las paradas de transporte público, según lo estipulado en el {info_proyecto.get('estudio', '[estudio]')} aprobado mediante Resolución Exenta {info_proyecto.get('resolucion', '[resolucion]')}, con fecha {info_proyecto.get('fecha', '[fecha]')}, en la comuna {info_proyecto.get('comuna', '[comuna]')}.")
        agregar_espacio(document)
        agregar_texto(document, f"Respecto a las medidas de mitigación mencionadas, se expone lo siguiente en el {info_proyecto.get('estudio', '[estudio]')} aprobado:")
        agregar_espacio(document)
        agregar_texto(document, info_proyecto.get('medida_mitigacion', '[medida_mitigacion]'))
        document.add_page_break()

        # ==========================================================
        # CAPÍTULO 2: DESCRIPCIÓN DEL PROYECTO
        # ==========================================================
        print("   - Creando Capítulo 2: Descripción del Proyecto...")
        cambiar_capitulo(estado_informe, 2)
        agregar_titulo(document, "2. DESCRIPCIÓN DEL PROYECTO")
        agregar_texto(document, f"El proyecto {info_proyecto.get('proyecto', '[nombre_proyecto]')}, se ubica en {info_proyecto.get('ubi_proyecto', '[ubi_proyecto]')}, comuna de {info_proyecto.get('comuna', '[comuna]')}, {info_proyecto.get('region', '[region]')}. En la siguiente figura N°2.1, se podrá visualizar la ubicación del proyecto:")
        agregar_espacio(document)
        if img_ubicacion_proyecto_id: agregar_imagen_con_formato_drive(document, service_drive, img_ubicacion_proyecto_id, descripcion="Ubicación del Proyecto", estado=estado_informe, fuente="Elaboración Propia en base a Google Earth")
        document.add_page_break()

        # ==========================================================
        # CAPÍTULO 3: INSPECCIÓN DE PARADEROS (VERSIÓN DETALLADA)
        # ==========================================================
        print("   - Creando Capítulo 3: Inspección de Paraderos...")
        cambiar_capitulo(estado_informe, 3)
        agregar_titulo(document, "3. INSPECCIÓN Y DESCRIPCIÓN DE PARADEROS INVOLUCRADOS")
        agregar_espacio(document)
        agregar_texto(document, "En este apartado se reporta la situación actual de las paradas en estudio, catastradas en las visitas a terreno. En la figura siguiente se muestra la ubicación actual de cada paradero:")
        if img_ubicacion_paradas_id:
            agregar_imagen_con_formato_drive(
                document, service_drive, img_ubicacion_paradas_id, 
                descripcion="Ubicación Paradas de Transporte Público en Estudio",
                estado=estado_informe,
                fuente="Elaboración Propia en base a Google Earth"
            )


        # --- Bucle para cada paradero (sub-capítulos de la sección 3) ---
        paraderos = datos_informe.get("paraderos", []) or []
        for i, paradero in enumerate(paraderos, start=1):
            document.add_page_break()

            info_paradero = paradero.get("info_paradero") or paradero.get("infoParadero") or {}
            # Fallbacks por si vinieran en otro nivel / nombres antiguos
            codigo = (
                info_paradero.get("codigo")
                or paradero.get("codigo")
                or paradero.get("codigo_paradero")
                or "S/C"
            )
            ubicacion = (
                info_paradero.get("ubicacion")
                or paradero.get("ubicacion")
                or paradero.get("ubicacion_paradero")
                or "Sin ubicación"
            )
            analisis = paradero.get("analisis") or {}

            print(f"   -> Procesando Paradero N°{i}: {codigo}")

            # Subtítulo
            agregar_subtitulo(document, f"3.{i} {codigo} - {ubicacion}")
            agregar_espacio(document)

            # Tabla 1: Imagen General
            crear_tabla_evidencia(document, service_drive, "Imagen general del paradero", analisis.get("general", {}))
            document.add_page_break()

            # Tabla 2: Refugio y Andén
            crear_tabla_evidencia(document, service_drive, "Evidencia Fotográfica de Refugio y Andén", analisis.get("refugio_anden", {}))
            document.add_page_break()

            # Tabla 3: Señal y Demarcación
            crear_tabla_evidencia(document, service_drive, "Evidencia Fotográfica de Señal y Demarcación", analisis.get("senal", {}))
            document.add_page_break()
            
            # Tabla 4: Características (desde los datos almacenados)
            filas = []
            for fila in paradero.get("tabla", []) or []:
                filas.append([
                    str(fila.get("caracteristica", "")).strip(),
                    str(fila.get("cumplimiento", "")).strip(),
                    str(fila.get("observacion", "")).strip(),
                ])

            if filas:
                headers = ["Característica", "Cumplimiento", "Observación"]
                agregar_tabla_formateada(
                    document,
                    headers=headers,
                    rows=filas,
                    titulo="Tabla de Características del Paradero",
                    estado=estado_informe,  # si tu función lo usa
                    fuente="Elaboración Propia"
                )


        # ==========================================================
        # CAPÍTULO 4: INFORMACIÓN DE PARADAS
        # ==========================================================
        print("   - Creando Capítulo 4: Información de Paradas...")
        document.add_page_break()
        cambiar_capitulo(estado_informe, 4)
        agregar_titulo(document, "4. INFORMACIÓN DE PARADAS DE TRANSPORTE PÚBLICO")
        agregar_espacio(document)
        agregar_texto(document, "En la siguiente tabla se reportan los servicios de bus que utilizan cada parada en estudio, con su respectivo destino:")
        agregar_espacio(document)

        if tablas_id:
            try:
                req = service_drive.files().get_media(fileId=tablas_id)  # IDs, no nombres
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, req)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                fh.seek(0)

                # Lee directamente desde el buffer en memoria
                df_resumen = pd.read_excel(fh, sheet_name="Paradas")  # ajusta nombre de hoja si corresponde
                agregar_tabla_formateada(document, df_resumen,
                                       "Información de Paradas de Transporte Públic",
                                       estado=estado_informe,
                                       fuente="Elaboración Propia en base DTPM - RED movilidad - Terreno")
            except Exception as e:
                agregar_texto(document, f"[ERROR: No se pudo leer 'Tablas.xlsx' desde Drive. Detalle: {e}]")
        else:
            agregar_texto(document, "[ERROR: No se recibió 'tablas_id' para cargar las tablas desde Drive.]")


        # ==========================================================
        # CAPÍTULO 5: MEDIDA DE MITIGACIÓN
        # ==========================================================
        print("   - Creando Capítulo 5: Medida de Mitigación...")
        document.add_page_break()
        cambiar_capitulo(estado_informe, 5)
        # Corregimos el error de tipeo de 'regar_titulo' a 'agregar_titulo'
        agregar_titulo(document, "5. MEDIDA DE MITIGACIÓN")
        agregar_espacio(document)
        agregar_texto(document, f"En función de la información recopilada en terreno, se presenta una tabla resumen con las mejoras a ejecutar de acuerdo con lo indicado en la aprobación del {info_proyecto.get('estudio', '[estudio]')} aprobado mediante {info_proyecto.get('resolucion', '[resolucion]')}. El cuadro que se presenta a continuación, indica un resumen con el estado de los paraderos revisados en la minuta y posteriormente se mencionan los elementos que requieren intervención.")
        agregar_espacio(document)

        if tablas_id:
            try:
                req = service_drive.files().get_media(fileId=tablas_id)  # IDs, no nombres
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, req)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                fh.seek(0)

                # Lee directamente desde el buffer en memoria
                df_resumen = pd.read_excel(fh, sheet_name="Resumen")  # ajusta nombre de hoja si corresponde
                agregar_tabla_formateada(document, df_resumen,
                                       "Resumen estado de Paraderos",
                                       estado=estado_informe,
                                       fuente="Elaboración Propia en base DTPM - RED movilidad - Terreno")
            except Exception as e:
                agregar_texto(document, f"[ERROR: No se pudo leer 'Tablas.xlsx' desde Drive. Detalle: {e}]")
        else:
            agregar_texto(document, "[ERROR: No se recibió 'tablas_id' para cargar las tablas desde Drive.]")
            
        # --- FINALIZACIÓN ---
        print("✅ Informe generado en memoria.")
        return document

    except Exception as e:
        print(f"❌ Error crítico durante la generación del informe: {e}")
        import traceback
        traceback.print_exc()
        return None